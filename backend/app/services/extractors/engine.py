import fitz  # PyMuPDF
import docx
from bs4 import BeautifulSoup
import os
from typing import Dict, Any, Tuple
import logging

from app.services.extractors.base import BaseExtractor
from app.models.knowledge import SourceType

logger = logging.getLogger(__name__)

class PDFExtractor(BaseExtractor):
    def extract(self, file_path: str) -> Tuple[str, int, Dict[str, Any]]:
        text_content = []
        metadata = {}
        page_count = 0
        try:
            with fitz.open(file_path) as doc:
                page_count = doc.page_count
                metadata = doc.metadata
                for page in doc:
                    text_content.append(page.get_text())
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {str(e)}")
            raise e
        return "\n\n".join(text_content), page_count, metadata

class DOCXExtractor(BaseExtractor):
    def extract(self, file_path: str) -> Tuple[str, int, Dict[str, Any]]:
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract paragraphs and headings
            for para in doc.paragraphs:
                if para.text.strip():
                    if para.style.name.startswith('Heading'):
                        text_content.append(f"## {para.text.strip()}")
                    else:
                        text_content.append(para.text.strip())
                        
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        text_content.append(" | ".join(row_data))
            
            metadata = {
                "author": doc.core_properties.author,
                "title": doc.core_properties.title,
                "created": str(doc.core_properties.created)
            }
            page_count = len(doc.sections) if doc.sections else 1
            return "\n\n".join(text_content), page_count, metadata
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {str(e)}")
            raise e

class HTMLExtractor(BaseExtractor):
    def extract(self, file_path: str) -> Tuple[str, int, Dict[str, Any]]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f, "html.parser")
                
            # Remove scripts, styles, and non-content tags
            for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
                element.decompose()
                
            metadata = {
                "title": soup.title.string if soup.title else None
            }
            text_content = soup.get_text(separator="\n\n")
            return text_content, 1, metadata
        except Exception as e:
            logger.error(f"Failed to parse HTML {file_path}: {str(e)}")
            raise e

class TXTExtractor(BaseExtractor):
    def extract(self, file_path: str) -> Tuple[str, int, Dict[str, Any]]:
        try:
            # Try UTF-8 first
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                encoding = "utf-8"
            except UnicodeDecodeError:
                # Fallback to ascii/latin-1
                with open(file_path, "r", encoding="latin-1") as f:
                    content = f.read()
                encoding = "latin-1"
                
            return content, 1, {"encoding": encoding}
        except Exception as e:
            logger.error(f"Failed to parse TXT {file_path}: {str(e)}")
            raise e

class MarkdownExtractor(BaseExtractor):
    def extract(self, file_path: str) -> Tuple[str, int, Dict[str, Any]]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                
            # Basic structural extraction (can be extended with markdown-it-py)
            lines = content.split('\n')
            headings = [line for line in lines if line.startswith('#')]
            code_blocks = content.count('```') // 2
            
            metadata = {
                "is_markdown": True,
                "headings_count": len(headings),
                "code_blocks_count": code_blocks
            }
            return content, 1, metadata
        except Exception as e:
            logger.error(f"Failed to parse Markdown {file_path}: {str(e)}")
            raise e

class ExtractionService:
    def __init__(self):
        self.extractors = {
            SourceType.PDF: PDFExtractor(),
            SourceType.DOCX: DOCXExtractor(),
            SourceType.HTML: HTMLExtractor(),
            SourceType.WEBSITE: HTMLExtractor(),
            SourceType.TXT: TXTExtractor(),
            SourceType.MARKDOWN: MarkdownExtractor(),
            SourceType.FAQ: TXTExtractor(),  # Just text
            SourceType.ARTICLE: HTMLExtractor()
        }

    def process_file(self, file_path: str, source_type: str) -> Tuple[str, int, Dict[str, Any]]:
        """
        Routes the file to the appropriate extractor based on source_type.
        If source_type doesn't perfectly map, it attempts to guess from extension.
        """
        extractor = self.extractors.get(source_type)
        
        if not extractor:
            # Guess from extension
            ext = os.path.splitext(file_path)[1].lower()
            if ext == ".pdf":
                extractor = self.extractors[SourceType.PDF]
            elif ext in [".docx", ".doc"]:
                extractor = self.extractors[SourceType.DOCX]
            elif ext in [".html", ".htm"]:
                extractor = self.extractors[SourceType.HTML]
            elif ext in [".md", ".markdown"]:
                extractor = self.extractors[SourceType.MARKDOWN]
            else:
                extractor = self.extractors[SourceType.TXT]

        return extractor.extract(file_path)

extraction_service = ExtractionService()
