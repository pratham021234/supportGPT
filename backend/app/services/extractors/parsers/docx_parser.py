import docx
import logging
from typing import Dict, Any, Tuple

from app.services.extractors.base import BaseExtractor

logger = logging.getLogger(__name__)

class DOCXExtractor(BaseExtractor):
    def extract(self, file_path: str) -> Tuple[str, int, Dict[str, Any]]:
        text_content = []
        metadata = {}
        pages_mapping = []
        
        try:
            doc = docx.Document(file_path)
            current_offset = 0
            
            # Since DOCX doesn't have a strict concept of physical pages like PDF,
            # we will treat sections or sequential chunks as page mappings.
            # But we can also just extract linearly.
            
            # Extract paragraphs and headings
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Check if heading
                    if para.style.name.startswith('Heading'):
                        # Determine heading level
                        try:
                            level = int(para.style.name.split(' ')[-1])
                            prefix = '#' * level
                            text = f"{prefix} {text}"
                        except:
                            text = f"## {text}"
                    # Check if list item
                    elif 'List' in para.style.name:
                        text = f"- {text}"
                        
                    start_offset = current_offset
                    end_offset = current_offset + len(text)
                    
                    text_content.append(text)
                    current_offset = end_offset + 2 # +2 for \n\n
                    
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    # Format as markdown table row
                    if row_data:
                        table_text.append("| " + " | ".join(row_data) + " |")
                
                if table_text:
                    # Add markdown table header separator
                    header_len = len(table.rows[0].cells) if table.rows else 0
                    if header_len > 0:
                        separator = "|" + "|".join(["---"] * header_len) + "|"
                        table_text.insert(1, separator)
                        
                    joined_table = "\n".join(table_text)
                    
                    start_offset = current_offset
                    end_offset = current_offset + len(joined_table)
                    
                    text_content.append(joined_table)
                    current_offset = end_offset + 2
                    
            # Extract Metadata
            props = doc.core_properties
            metadata = {
                "author": props.author or "",
                "title": props.title or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "category": props.category or "",
                "comments": props.comments or "",
            }
            
            page_count = len(doc.sections) if doc.sections else 1
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {str(e)}")
            raise e
            
        return "\n\n".join(text_content), page_count, metadata
